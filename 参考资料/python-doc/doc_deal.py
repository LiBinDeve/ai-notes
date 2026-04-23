
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches

from matplotlib.figure import Figure
from docx2pdf import convert
import io

# docx默认只能识别英文字体，中文字体需设置成这个qn('w:eastAsia')
def set_chinese_font_for_style(style, font_name='仿宋'):
    rPr = style._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)

# =========自定义字体========
def ensure_styles(doc):

    # 仿宋 三号
    style_main = doc.styles.add_style('MyMainStyle', WD_STYLE_TYPE.PARAGRAPH)
    style_main.base_style = doc.styles['Normal']
    style_main.font.name = '仿宋'
    style_main.font.size = Pt(16)
    style_main.font.bold = False
    set_chinese_font_for_style(style_main, '仿宋')

    # 仿宋 小二
    style_heading = doc.styles.add_style('MyHeadingStyle', WD_STYLE_TYPE.PARAGRAPH)
    style_heading.base_style = doc.styles['Normal']
    style_heading.font.name = '仿宋'
    style_heading.font.size = Pt(18)
    style_heading.font.bold = True
    set_chinese_font_for_style(style_heading, '仿宋')

    # 宋体 五号
    style_heading = doc.styles.add_style('MyGraghHeadingStyle', WD_STYLE_TYPE.PARAGRAPH)
    style_heading.base_style = doc.styles['Normal']
    style_heading.font.name = '宋体'
    style_heading.font.size = Pt(10.5)
    style_heading.font.bold = True
    set_chinese_font_for_style(style_heading, '仿宋')

    # 方正小标宋 简体 标题
    style_title = doc.styles.add_style('MyTitleStyle', WD_STYLE_TYPE.PARAGRAPH)
    style_title.base_style = doc.styles['Normal']
    style_title.font.name = '方正小标宋简体'
    style_title.font.size = Pt(22)
    style_title.font.bold = False
    style_title.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    set_chinese_font_for_style(style_title, '方正小标宋简体')

# 创建文档
def create_doc():
    doc = Document()
    return doc

# 写入提示
def write_hint(doc: Document, text: str):
    para_hint = doc.add_paragraph(text)
    para_hint.style = 'MyMainStyle'
    para_hint.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para_hint.paragraph_format.right_indent = Pt(28.2)

# 写入主标题
def write_title(doc: Document, month: int):
    headingmain = doc.add_paragraph(f'商州区突发公共卫生事件及需关注的\n传染病{month}月风险评估报告')
    headingmain.style = 'MyTitleStyle'
    headingmain.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 写入标注1
def write_note_one(doc: Document, year:int, month: int):
    note_one = doc.add_paragraph()
    note_one.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_one_run = note_one.add_run(f'{year}年第{month}期')
    note_one_run.font.name = '仿宋'
    note_one_run.font.size = Pt(18)
    note_one_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

# 写入标注2
def write_note_two(doc: Document, year:int, month: int):
    note_two = doc.add_paragraph()
    note_two_run = note_two.add_run(f'\t商州区疾病预防控制中心\t\t\t{year}年{month}月1日\t')
    note_two_run.font.name = '仿宋'
    note_two_run.font.size = Pt(16)
    note_two_run.font.underline = True
    note_two_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

# 写入一级小标题
def write_level_one_head(doc: Document, text: str):
    heading1 = doc.add_paragraph(text)
    heading1.style = 'MyHeadingStyle'
    heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading1.paragraph_format.left_indent = Pt(28.2)

# 写入二级小标题
def write_level_two_head(doc: Document, text: str, number: str):
    heading = doc.add_paragraph(f'（{number}）{text}')
    heading.style = 'MyMainStyle'
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading.paragraph_format.left_indent = Pt(28.2)   

# 写入正文
def write_text(doc: Document, text: str):
    text = doc.add_paragraph(text)
    text.style = 'MyMainStyle'
    text.paragraph_format.first_line_indent = Pt(28.2)
    text.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    text.paragraph_format.line_spacing = 1.1

# 写入图
def write_graph(doc: Document, fig: Figure, width=Inches(6)):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=300, bbox_inches="tight")
    buf.seek(0)  
    doc.add_picture(buf, width=width)
    buf.close()

# 写入图标题
def write_graph_head(doc: Document, disease: str, number:int):
    graph_heading = doc.add_paragraph(f'图{number} 商州区{disease}报告发病分布情况')
    graph_heading.style = 'MyGraghHeadingStyle'
    graph_heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER 

# 写入空行
def write_blank_line(doc: Document):
    doc.add_paragraph()

# 保存文档
import pythoncom
from docx2pdf import convert

def save_doc_as_pdf(doc, docx_path: str, pdf_path: str) -> bytes:
    # 先保存 DOCX
    doc.save(docx_path)

    # 关键：初始化 COM，并让 docx2pdf 复用 Word 实例，降低崩溃概率
    pythoncom.CoInitialize()
    try:
        # keep_active=True 让 docx2pdf 复用同一个 Word，而不是每次启动/退出
        convert(docx_path, pdf_path, keep_active=True)
    finally:
        pythoncom.CoUninitialize()

    with open(pdf_path, "rb") as f:
        return f.read()

# 读取文档
def read_docx_text(file_path: str) -> str:

    doc = Document(file_path)
    # 提取非空段落
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # 拼成一个完整字符串
    return "\n".join(paragraphs)

